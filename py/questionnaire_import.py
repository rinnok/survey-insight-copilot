"""Questionnaire file extraction and review helpers.

The parser deliberately produces a review draft. It never treats extracted
questions as confirmed until a user has compared them with the source file.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".xlsx", ".xlsm", ".csv", ".docx", ".pdf", ".txt"}
QUESTION_RE = re.compile(
    r"^\s*(?P<id>(?:Q|Ｑ|問)\s*\d+(?:[-_.]\d+)?|\d+(?:[-_.]\d+)?)"
    r"[\s　:：.)）-]*(?P<text>.+?)\s*$",
    re.IGNORECASE,
)
OPTION_RE = re.compile(r"(?:^|[\s　])(?:\d+|[①-⑳]|[ア-ン])[.)）:：、]\s*")


def _clean(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    return re.sub(r"[ \t　]+", " ", str(value).replace("\r", "\n")).strip()


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("テキストの文字コードを判定できませんでした")


def _extract_text(filename: str, raw: bytes) -> tuple[str, list[str]]:
    suffix = Path(filename).suffix.lower()
    warnings: list[str] = []
    if suffix == ".txt":
        return _decode_text(raw), warnings
    if suffix == ".docx":
        document = Document(io.BytesIO(raw))
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(blocks), warnings
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        if len(text) < 30:
            raise ValueError("PDFから文字を抽出できませんでした。画像スキャンPDFは、OCR済みPDFにして再度添付してください")
        if any(not page for page in pages):
            warnings.append("文字を抽出できないページがあります。原本のページ数と照合してください")
        return text, warnings
    raise ValueError("この形式は文章抽出に対応していません")


def _find_column(columns: list[str], candidates: tuple[str, ...]) -> str:
    normalized = {re.sub(r"\s+", "", col).lower(): col for col in columns}
    for candidate in candidates:
        for key, original in normalized.items():
            if candidate in key:
                return original
    return ""


def _infer_type(text: str, options: list[str], declared: str = "") -> str:
    value = declared.strip()
    if value:
        return value
    joined = f"{text} {' '.join(options)}"
    if re.search(r"複数|あてはまるものすべて|いくつでも", joined):
        return "複数選択"
    if re.search(r"自由記述|具体的に|理由をお書き|その他.*記入", joined):
        return "自由記述"
    if re.search(r"\d+\s*段階|まったく.*非常に|そう思わない.*そう思う", joined):
        return "尺度"
    if options:
        return "単一選択"
    return "未判定"


def _split_options(value: str) -> list[str]:
    value = _clean(value)
    if not value:
        return []
    parts = re.split(r"\s*(?:／|/|\||;|；|\n|、(?=\S{1,18}(?:、|$)))\s*", value)
    return [part.strip(" ・-") for part in parts if part.strip(" ・-")]


def _questions_from_table(frame: pd.DataFrame) -> tuple[list[dict[str, Any]], str, list[str]]:
    frame = frame.dropna(how="all").fillna("")
    if frame.empty:
        return [], "", ["表にデータがありません"]
    frame.columns = [_clean(col) or f"列{i + 1}" for i, col in enumerate(frame.columns)]
    columns = list(frame.columns)
    id_col = _find_column(columns, ("設問番号", "質問番号", "qid", "id", "no", "番号"))
    text_col = _find_column(columns, ("設問文", "質問文", "設問", "質問", "項目"))
    type_col = _find_column(columns, ("回答形式", "設問形式", "形式", "タイプ", "type"))
    option_col = _find_column(columns, ("選択肢", "回答項目", "尺度", "options", "option"))
    warnings: list[str] = []

    if not text_col:
        # A common simple questionnaire format is Q number in the first column
        # and question text in the second column.
        text_col = columns[1] if len(columns) > 1 else columns[0]
        id_col = id_col or (columns[0] if len(columns) > 1 else "")
        warnings.append("列見出しから設問文を特定できなかったため、表の先頭列をもとに推定しました")

    questions: list[dict[str, Any]] = []
    source_lines = ["\t".join(columns)]
    for index, row in frame.iterrows():
        values = [_clean(row[col]) for col in columns]
        source_lines.append("\t".join(values))
        text = _clean(row[text_col])
        if not text:
            continue
        qid = _clean(row[id_col]) if id_col else f"Q{len(questions) + 1}"
        options = _split_options(_clean(row[option_col])) if option_col else []
        declared = _clean(row[type_col]) if type_col else ""
        questions.append({
            "id": qid or f"Q{len(questions) + 1}",
            "text": text,
            "type": _infer_type(text, options, declared),
            "options": options,
            "source": f"表の{index + 2}行目",
        })
    return questions, "\n".join(source_lines), warnings


def _questions_from_text(text: str) -> tuple[list[dict[str, Any]], list[str]]:
    lines = [_clean(line) for line in text.splitlines()]
    questions: list[dict[str, Any]] = []
    warnings: list[str] = []
    current: dict[str, Any] | None = None

    for line_number, line in enumerate(lines, 1):
        if not line:
            continue
        match = QUESTION_RE.match(line)
        if (
            current
            and match
            and re.match(r"^\d", match.group("id"))
            and re.match(r"^(?:Q|Ｑ|問)", current["id"], re.IGNORECASE)
        ):
            current["options"].append(match.group("text").strip())
            continue
        if match:
            if current:
                questions.append(current)
            current = {
                "id": re.sub(r"\s+", "", match.group("id")).upper(),
                "text": match.group("text").strip(),
                "type": "未判定",
                "options": [],
                "source": f"抽出テキスト {line_number}行目",
            }
            continue
        if not current:
            continue
        if OPTION_RE.search(line) or re.match(r"^[□○●✓✔]", line):
            parts = [part.strip(" □○●✓✔・") for part in OPTION_RE.split(line) if part.strip(" □○●✓✔・")]
            current["options"].extend(parts or [line])
        elif len(line) <= 120 and current["type"] == "未判定":
            current["text"] = f"{current['text']} {line}".strip()

    if current:
        questions.append(current)
    for question in questions:
        question["options"] = list(dict.fromkeys(question["options"]))
        question["type"] = _infer_type(question["text"], question["options"])
    if not questions:
        warnings.append("設問番号（例：Q1、問1、1.）を検出できませんでした")
    return questions, warnings


def parse_questionnaire(filename: str, raw: bytes) -> dict[str, Any]:
    safe_name = Path(filename).name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError("対応形式は Excel / CSV / Word / PDF / TXT です")
    if not raw:
        raise ValueError("ファイルが空です")

    warnings: list[str] = []
    if suffix in {".xlsx", ".xlsm", ".csv"}:
        if suffix == ".csv":
            frame = None
            for encoding in ("utf-8-sig", "utf-8", "cp932"):
                try:
                    frame = pd.read_csv(io.BytesIO(raw), encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if frame is None:
                raise ValueError("CSVの文字コードを判定できませんでした")
        else:
            frame = pd.read_excel(io.BytesIO(raw))
        questions, source_text, table_warnings = _questions_from_table(frame)
        warnings.extend(table_warnings)
    else:
        source_text, extract_warnings = _extract_text(safe_name, raw)
        questions, text_warnings = _questions_from_text(source_text)
        warnings.extend(extract_warnings + text_warnings)

    if not questions:
        raise ValueError("設問を読み取れませんでした。設問番号と設問文が分かる形式に整えて再度添付してください")
    undecided = sum(1 for question in questions if question["type"] == "未判定")
    duplicate_ids = len(questions) - len({question["id"] for question in questions})
    if undecided:
        warnings.append(f"回答形式を判定できない設問が{undecided}件あります")
    if duplicate_ids:
        warnings.append(f"設問番号の重複が{duplicate_ids}件あります")

    return {
        "filename": safe_name,
        "questions": questions,
        "sourceText": source_text[:60000],
        "sourceTruncated": len(source_text) > 60000,
        "warnings": warnings,
        "stats": {
            "questionCount": len(questions),
            "undecidedCount": undecided,
            "duplicateIdCount": duplicate_ids,
        },
    }


def validate_confirmed_questions(questions: Any) -> list[dict[str, Any]]:
    if not isinstance(questions, list) or not questions:
        raise ValueError("確認済みの設問がありません")
    clean_questions: list[dict[str, Any]] = []
    for index, item in enumerate(questions, 1):
        if not isinstance(item, dict):
            raise ValueError(f"{index}件目の設問形式が不正です")
        qid = _clean(item.get("id")) or f"Q{index}"
        text = _clean(item.get("text"))
        qtype = _clean(item.get("type"))
        options = [_clean(option) for option in item.get("options", []) if _clean(option)]
        if not text:
            raise ValueError(f"{qid}の設問文を入力してください")
        if qtype == "未判定" or not qtype:
            raise ValueError(f"{qid}の回答形式を確認してください")
        clean_questions.append({"id": qid, "text": text, "type": qtype, "options": options})
    if len({item["id"] for item in clean_questions}) != len(clean_questions):
        raise ValueError("設問番号が重複しています")
    return clean_questions
